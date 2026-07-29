"""
Role Scout resume export (added Jul 2026, verification link build)
==================================================================
POST /export/resume        -> renders a verified tailored resume as a
                              downloadable PDF or Word document (binary)
GET  /export/resume/health -> reports pdf and docx renderer availability

Contract (JSON body):
{
  "format": "pdf" | "docx",
  "seeker_name": "Full Name",
  "contact_lines": ["optional strings under the name"],
  "posting": {"title": "...", "employer": "...", "location": "..."},
  "sections": [{"title": "SUMMARY", "lines": ["...", "..."]}],
  "verified_at": "2026-07-28",       optional
  "attested_version": 2              optional
}

sections is the only required content field. When verified_at and
attested_version are both present, a single quiet verification footer
line renders at the end of the document; when either is absent the
document renders with no footer, so the caller controls inclusion.

Output is deliberately ATS shaped: one column, standard fonts, plain
uppercase section headings, no tables, no graphics. The content arrives
already verified by the tailor lane, so this module renders mechanically
and applies no content gates.

CORS: the Role Scout app in the browser calls this route directly, so
the blueprint answers OPTIONS preflights and sends open CORS headers on
its responses. The route holds no secrets and writes nothing.

Lazy imports per the resume_extract pattern: a missing dependency
degrades to a 503 on this route instead of breaking the app at import.
Stateless, no worker memory.
"""

from flask import Blueprint, request, Response, send_file
from io import BytesIO
import json as _json
import re as _re

export_bp = Blueprint("resume_export", __name__)

_CORS_HEADERS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "POST, GET, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
    "Access-Control-Max-Age": "86400",
}


@export_bp.after_request
def _add_cors(resp):
    for k, v in _CORS_HEADERS.items():
        resp.headers[k] = v
    return resp


def _pdf_ready():
    try:
        import reportlab  # noqa: F401
        return True
    except Exception:
        return False


def _docx_ready():
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def _err(msg, status):
    return Response(_json.dumps({"error": msg}), status=status,
                    mimetype="application/json")


def _safe_name(*parts):
    """Builds a filesystem safe download name from name and posting parts."""
    joined = "_".join(p for p in parts if p)
    cleaned = _re.sub(r"[^A-Za-z0-9]+", "_", joined).strip("_")
    return cleaned or "Tailored_Resume"


def _read_payload(req):
    data = req.get_json(silent=True)
    if not isinstance(data, dict):
        raw = req.get_data(as_text=True) or ""
        try:
            data = _json.loads(raw)
        except Exception:
            data = None
    if not isinstance(data, dict):
        return None, _err("JSON body required", 400)

    sections = data.get("sections")
    if not isinstance(sections, list) or not sections:
        return None, _err("sections is required and must be a non empty list", 422)
    clean_sections = []
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        title = str(sec.get("title") or "").strip()
        lines = [str(x).strip() for x in (sec.get("lines") or []) if str(x).strip()]
        if lines:
            clean_sections.append({"title": title, "lines": lines})
    if not clean_sections:
        return None, _err("sections contained no usable lines", 422)

    posting = data.get("posting") if isinstance(data.get("posting"), dict) else {}
    payload = {
        "format": str(data.get("format") or "pdf").strip().lower(),
        "seeker_name": str(data.get("seeker_name") or "").strip(),
        "contact_lines": [str(x).strip() for x in (data.get("contact_lines") or [])
                          if str(x).strip()],
        "posting_title": str(posting.get("title") or "").strip(),
        "posting_employer": str(posting.get("employer") or "").strip(),
        "posting_location": str(posting.get("location") or "").strip(),
        "sections": clean_sections,
        "verified_at": str(data.get("verified_at") or "").strip(),
        "attested_version": data.get("attested_version"),
    }
    if payload["format"] not in ("pdf", "docx"):
        return None, _err("format must be pdf or docx", 422)
    return payload, None


_FOOTER_TEMPLATE = ("Verified by Role Scout against Version {v} of the attested "
                    "source resume on {d}. Every line traces to attested content. "
                    "Nothing was invented.")


def _footer_line(payload):
    """Returns the verification footer, or None when either fact is absent."""
    v = payload.get("attested_version")
    d = payload.get("verified_at")
    if v in (None, "") or not d:
        return None
    return _FOOTER_TEMPLATE.format(v=v, d=d)


def _render_pdf(payload):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from xml.sax.saxutils import escape

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                            topMargin=0.8 * inch, bottomMargin=0.8 * inch,
                            title=payload["seeker_name"] or "Tailored Resume")
    name_style = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=16,
                                leading=20, alignment=TA_LEFT)
    contact_style = ParagraphStyle("contact", fontName="Helvetica", fontSize=9.5,
                                   leading=13)
    posting_style = ParagraphStyle("posting", fontName="Helvetica-Oblique",
                                   fontSize=9.5, leading=13)
    heading_style = ParagraphStyle("heading", fontName="Helvetica-Bold",
                                   fontSize=11, leading=15, spaceBefore=10,
                                   spaceAfter=3)
    body_style = ParagraphStyle("body", fontName="Helvetica", fontSize=10,
                                leading=14, spaceAfter=2)
    footer_style = ParagraphStyle("footer", fontName="Helvetica", fontSize=8,
                                  leading=11, textColor="#555555", spaceBefore=14)

    story = []
    if payload["seeker_name"]:
        story.append(Paragraph(escape(payload["seeker_name"]), name_style))
    for line in payload["contact_lines"]:
        story.append(Paragraph(escape(line), contact_style))
    posting_bits = [payload["posting_title"], payload["posting_employer"],
                    payload["posting_location"]]
    posting_line = ", ".join(b for b in posting_bits if b)
    if posting_line:
        story.append(Spacer(1, 4))
        story.append(Paragraph(escape(posting_line), posting_style))
    for sec in payload["sections"]:
        if sec["title"]:
            story.append(Paragraph(escape(sec["title"].upper()), heading_style))
        for line in sec["lines"]:
            story.append(Paragraph(escape(line), body_style))
    footer = _footer_line(payload)
    if footer:
        story.append(Paragraph(escape(footer), footer_style))
    doc.build(story)
    buf.seek(0)
    return buf


def _render_docx(payload):
    from docx import Document
    from docx.shared import Pt, RGBColor

    d = Document()
    if payload["seeker_name"]:
        p = d.add_paragraph()
        run = p.add_run(payload["seeker_name"])
        run.bold = True
        run.font.size = Pt(16)
    for line in payload["contact_lines"]:
        p = d.add_paragraph()
        p.add_run(line).font.size = Pt(9.5)
    posting_bits = [payload["posting_title"], payload["posting_employer"],
                    payload["posting_location"]]
    posting_line = ", ".join(b for b in posting_bits if b)
    if posting_line:
        p = d.add_paragraph()
        run = p.add_run(posting_line)
        run.italic = True
        run.font.size = Pt(9.5)
    for sec in payload["sections"]:
        if sec["title"]:
            p = d.add_paragraph()
            run = p.add_run(sec["title"].upper())
            run.bold = True
            run.font.size = Pt(11)
        for line in sec["lines"]:
            p = d.add_paragraph()
            p.add_run(line).font.size = Pt(10)
    footer = _footer_line(payload)
    if footer:
        p = d.add_paragraph()
        run = p.add_run(footer)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@export_bp.route("/export/resume/health", methods=["GET"])
def export_health():
    body = {"ok": True, "pdf": _pdf_ready(), "docx": _docx_ready(),
            "version": "v1.0"}
    return Response(_json.dumps(body), mimetype="application/json")


@export_bp.route("/export/resume", methods=["POST", "OPTIONS"])
def export_resume():
    if request.method == "OPTIONS":
        return Response(status=204)
    payload, err = _read_payload(request)
    if err is not None:
        return err
    fmt = payload["format"]
    if fmt == "pdf":
        if not _pdf_ready():
            return _err("pdf renderer unavailable on this deploy", 503)
        buf = _render_pdf(payload)
        ext, mime = "pdf", "application/pdf"
    else:
        if not _docx_ready():
            return _err("docx renderer unavailable on this deploy", 503)
        buf = _render_docx(payload)
        ext, mime = "docx", ("application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document")
    fname = _safe_name(payload["seeker_name"], payload["posting_title"])
    return send_file(buf, mimetype=mime, as_attachment=True,
                     download_name="{0}.{1}".format(fname, ext))
