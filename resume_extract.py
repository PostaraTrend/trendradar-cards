"""
Role Scout resume extraction endpoint (added Jul 2026, RS PL 001)
================================================================
POST /extract          -> raw docx or pdf bytes in, JSON text out
GET  /extract/health   -> dependency and version report

Purpose: the parse lane in n8n downloads a stored resume file from
Supabase storage and posts the raw bytes here. This module performs the
mechanical step only: bytes to plain text plus page facts. It never
interprets, never summarizes, never reorders. The transcription model in
the lane does the line and section work afterwards, per RS PL 001.

Detection is by magic bytes, not filename: PK.. means docx (zip
container), %PDF means pdf. Anything else returns 415.

Response shape on success:
  {ok: true, kind: "docx"|"pdf", pages: int, chars: int,
   scanned: bool, truncated: bool, text: "..."}
pages for a pdf is the true page count; for a docx it is an estimate
from character volume (about 1800 characters per page), minimum 1.
scanned is true when a pdf has pages but effectively no text layer
(under 40 characters per page on average). The lane turns scanned into
parse_failed with a plain message, per the approved design: OCR is
deferred, honesty over guessing.

House rules honoured: no import time downloads or raises (lazy imports,
missing dependencies reported by health, never fatal to the app), no
worker memory, stateless, single request in and out.
"""

from flask import Blueprint, request, jsonify

extract_bp = Blueprint("extract", __name__)

EXTRACT_VERSION = "1.0"
TEXT_CHAR_CAP = 250000
DOCX_CHARS_PER_PAGE = 1800
SCANNED_CHARS_PER_PAGE = 40


def _docx_available():
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def _pdf_available():
    try:
        import pypdf  # noqa: F401
        return True
    except Exception:
        return False


def _extract_docx(data):
    """Paragraphs in document order, then table cells row by row.
    Verbatim text only. Empty paragraphs are dropped because they are
    layout, not content."""
    from io import BytesIO
    from docx import Document
    doc = Document(BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append("  ".join(cells))
    text = "\n".join(parts)
    pages = max(1, (len(text) + DOCX_CHARS_PER_PAGE - 1) // DOCX_CHARS_PER_PAGE)
    return text, pages


def _extract_pdf(data):
    from io import BytesIO
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(data))
    pages = len(reader.pages)
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            chunks.append("")
    text = "\n".join(chunks).strip()
    return text, pages


@extract_bp.get("/extract/health")
def extract_health():
    docx_ok = _docx_available()
    pdf_ok = _pdf_available()
    return jsonify({
        "ok": docx_ok and pdf_ok,
        "version": EXTRACT_VERSION,
        "docx": docx_ok,
        "pdf": pdf_ok,
        "text_char_cap": TEXT_CHAR_CAP,
    })


@extract_bp.post("/extract")
def extract():
    data = request.get_data()
    if not data or len(data) < 8:
        return jsonify({"ok": False, "error": "file body required"}), 400

    if data[:4] == b"PK\x03\x04":
        kind = "docx"
    elif data[:5] == b"%PDF-":
        kind = "pdf"
    else:
        return jsonify({"ok": False,
                        "error": "unsupported file type, docx or pdf only"}), 415

    try:
        if kind == "docx":
            if not _docx_available():
                return jsonify({"ok": False,
                                "error": "docx support not installed"}), 503
            text, pages = _extract_docx(data)
        else:
            if not _pdf_available():
                return jsonify({"ok": False,
                                "error": "pdf support not installed"}), 503
            text, pages = _extract_pdf(data)
    except Exception as exc:  # unreadable or corrupt file
        return jsonify({"ok": False,
                        "error": "file could not be read: %s" % exc.__class__.__name__}), 422

    scanned = False
    if kind == "pdf" and pages > 0:
        scanned = (len(text) / pages) < SCANNED_CHARS_PER_PAGE

    truncated = len(text) > TEXT_CHAR_CAP
    if truncated:
        text = text[:TEXT_CHAR_CAP]

    return jsonify({
        "ok": True,
        "kind": kind,
        "pages": pages,
        "chars": len(text),
        "scanned": scanned,
        "truncated": truncated,
        "text": text,
    })
